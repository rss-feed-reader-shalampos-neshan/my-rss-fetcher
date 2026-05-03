import feedparser
import requests
from newspaper import Article
import os
import json
import re
import time
from datetime import datetime  # اضافه شدن برای ثبت تاریخ در نام فایل
from docx import Document
from docx.shared import Pt

with open('feeds.txt', 'r', encoding='utf-8') as f:
    feeds = [line.strip() for line in f if line.strip()]

PROCESSED_FILE = 'processed_urls.json'
if os.path.exists(PROCESSED_FILE):
    with open(PROCESSED_FILE, 'r', encoding='utf-8') as f:
        processed_urls = set(json.load(f))
else:
    processed_urls = set()

new_processed = set()
os.makedirs('articles', exist_ok=True)

# User-Agent مثل یک مرورگر واقعی
headers = {
    'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36'
}

for idx, feed_url in enumerate(feeds):
    print(f"در حال بررسی فید: {feed_url}")
    try:
        feed = feedparser.parse(feed_url)
    except Exception as e:
        print(f"خطا در فید {feed_url}: {e}")
        continue

    for entry in feed.entries:
        article_url = entry.link

        # پرش از لینک‌های کامنت یا غیرمفید
        if "#comment" in article_url or article_url in processed_urls:
            continue

        print(f"مقاله جدید یافت شد: {article_url}")
        try:
            # تنظیم هدر برای کتابخانه newspaper
            article = Article(article_url, language='en')
            article.config.browser_user_agent = headers['User-Agent']
            article.download()
            article.parse()

            title = article.title or "بدون عنوان"
            authors = ', '.join(article.authors) if article.authors else 'ناشناس'
            publish_date = article.publish_date.strftime('%Y-%m-%d %H:%M') if article.publish_date else 'نامشخص'
            text_content = article.text

            if not text_content:
                print(f"متنی پیدا نشد: {article_url}")
                continue

            doc = Document()
            style = doc.styles['Normal']
            font = style.font
            font.name = 'Calibri'
            font.size = Pt(12)

            doc.add_heading(title, level=1)
            doc.add_paragraph(f"منبع: {article_url}")
            doc.add_paragraph(f"نویسنده: {authors}")
            doc.add_paragraph(f"تاریخ انتشار: {publish_date}")
            doc.add_paragraph("").add_run("─" * 50)

            for paragraph_text in text_content.split('\n'):
                if paragraph_text.strip():
                    doc.add_paragraph(paragraph_text.strip())

            safe_title = re.sub(r'[<>:"/\\|?*]', '_', title)[:80]
            timestamp = datetime.utcnow().strftime('%Y-%m-%d_%H%M')  # زمان واکشی مقاله
            filename = f"articles/{timestamp}_{safe_title}_{idx}.docx"
            doc.save(filename)
            print(f"ذخیره شد: {filename}")

            processed_urls.add(article_url)
            new_processed.add(article_url)

            # مکث بیشتر برای احترام به سرور (مخصوصاً برای سایت‌های حساس مثل CBR)
            time.sleep(5)

        except Exception as e:
            print(f"خطا در استخراج {article_url}: {e}")
            continue

if new_processed:
    with open(PROCESSED_FILE, 'w', encoding='utf-8') as f:
        json.dump(list(processed_urls), f, indent=2)
    print(f"{len(new_processed)} مقاله جدید اضافه شد.")
else:
    print("هیچ مقالۀ جدیدی یافت نشد.")
