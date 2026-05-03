import requests
import os
import json
import re
import time
from datetime import datetime, timezone   # تغییر: timezone اضافه شد
from urllib.parse import urlparse, parse_qs
from newspaper import Article
from docx import Document
from docx.shared import Pt

with open('manual_urls.txt', 'r', encoding='utf-8') as f:
    urls = [line.strip() for line in f if line.strip()]

PROCESSED_FILE = 'processed_manual_urls.json'
if os.path.exists(PROCESSED_FILE):
    with open(PROCESSED_FILE, 'r', encoding='utf-8') as f:
        processed_urls = set(json.load(f))
else:
    processed_urls = set()

new_processed = set()
os.makedirs('manual_articles', exist_ok=True)

headers = {
    'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36'
}

for idx, raw_url in enumerate(urls):
    article_url = raw_url

    if 'google.com/url' in article_url:
        parsed = urlparse(article_url)
        query_params = parse_qs(parsed.query)
        if 'url' in query_params:
            article_url = query_params['url'][0]
            print(f"لینک گوگل شناسایی شد. استفاده از لینک اصلی: {article_url}")

    if article_url in processed_urls:
        print(f"مقاله قبلاً ذخیره شده: {article_url}")
        continue

    print(f"دریافت مقاله: {article_url}")
    try:
        article = Article(article_url, language='en')
        article.config.browser_user_agent = headers['User-Agent']
        article.download()
        article.parse()

        title = article.title or "بدون عنوان"
        authors = ', '.join(article.authors) if article.authors else 'ناشناس'
        publish_date = article.publish_date.strftime('%Y-%m-%d %H:%M') if article.publish_date else 'نامشخص'
        text_content = article.text

        if not text_content:
            print(f"متنی پیدا نشد: {article_url}")
            continue

        doc = Document()
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Calibri'
        font.size = Pt(12)

        doc.add_heading(title, level=1)
        doc.add_paragraph(f"منبع: {article_url}")
        doc.add_paragraph(f"نویسنده: {authors}")
        doc.add_paragraph(f"تاریخ انتشار: {publish_date}")
        doc.add_paragraph("").add_run("─" * 50)

        for paragraph_text in text_content.split('\n'):
            if paragraph_text.strip():
                doc.add_paragraph(paragraph_text.strip())

        safe_title = re.sub(r'[<>:"/\\|?*]', '_', title)[:80]
        # تغییر: استفاده از timezone.utc به‌جای datetime.UTC
        timestamp = datetime.now(timezone.utc).strftime('%Y-%m-%d_%H%M')
        filename = f"manual_articles/{timestamp}_{safe_title}.docx"
        doc.save(filename)
        print(f"ذخیره شد: {filename}")

        processed_urls.add(article_url)
        new_processed.add(article_url)

        time.sleep(5)

    except Exception as e:
        print(f"خطا در استخراج {article_url}: {e}")
        continue

if new_processed:
    with open(PROCESSED_FILE, 'w', encoding='utf-8') as f:
        json.dump(list(processed_urls), f, indent=2)
    print(f"{len(new_processed)} مقاله جدید اضافه شد.")
else:
    print("هیچ مقالۀ جدیدی یافت نشد.")

with open('manual_urls.txt', 'w', encoding='utf-8') as f:
    f.write('')
print("فایل manual_urls.txt خالی شد.")
